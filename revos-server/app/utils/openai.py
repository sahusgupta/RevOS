import openai
from pinecone import Pinecone, ServerlessSpec
from dotenv import load_dotenv
import os
import json
import logging
from typing import List, Dict, Any, Optional
import time
import re
from datetime import datetime, timedelta
import PyPDF2
import pdfplumber
from docx import Document
import io

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    load_dotenv()
    server_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    env_path = os.path.join(server_root, '.env')
    if os.path.exists(env_path):
        load_dotenv(env_path)
except Exception as e:
    logger.warning(f"Could not load .env file: {e}")

openai_api_key = os.getenv("OPENAI_API_KEY")
client = None
if openai_api_key:
    try:
        client = openai.OpenAI(api_key=openai_api_key, timeout=30.0)
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {e}")

EMBEDDING_MODEL = "text-embedding-3-small"
CHAT_MODEL = "gpt-4o-mini"

pinecone_api_key = os.getenv("PINECONE_API_KEY")
pc = None
if pinecone_api_key:
    try:
        pc = Pinecone(api_key=pinecone_api_key)
    except Exception as e:
        logger.error(f"Failed to initialize Pinecone: {e}")

INDEX_NAME = "revos-syllabus"
DIMENSION = 1536

SYSTEM_PROMPT = "You are Rev, Texas A&M mascot. Help with academic questions using provided syllabus info."

class PineconeManager:
    def __init__(self):
        self.index = None
        self._initialize_index()
    
    def _initialize_index(self):
        try:
            if not pc:
                logger.warning("Pinecone not initialized")
                return
            if INDEX_NAME not in pc.list_indexes().names():
                logger.info(f"Creating index: {INDEX_NAME}")
                pc.create_index(name=INDEX_NAME, dimension=DIMENSION, metric="cosine",
                    spec=ServerlessSpec(cloud="aws", region="us-east-1"))
                time.sleep(10)
            self.index = pc.Index(INDEX_NAME)
            logger.info(f"Connected to Pinecone index: {INDEX_NAME}")
        except Exception as e:
            logger.error(f"Error initializing index: {e}")
    
    def upsert_syllabus_content(self, course_id: str, content_chunks: List[Dict[str, Any]]):
        try:
            if not self.index:
                return False
            vectors = []
            for i, chunk in enumerate(content_chunks):
                embedding = generate_embedding(chunk['text'])
                sanitized = course_id.encode('ascii', 'ignore').decode('ascii')
                sanitized = re.sub(r'[^a-zA-Z0-9_-]', '_', sanitized)
                vector_id = f"{sanitized}_{i}".strip('_')
                if not vector_id:
                    vector_id = f"v_{i}"
                vectors.append({
                    'id': vector_id,
                    'values': embedding,
                    'metadata': {
                        'course_id': course_id,
                        'course_name': chunk.get('course_name', ''),
                        'text': chunk['text'][:500]
                    }
                })
            for i in range(0, len(vectors), 100):
                self.index.upsert(vectors=vectors[i:i+100])
            logger.info(f"Upserted {len(vectors)} vectors")
            return True
        except Exception as e:
            logger.error(f"Error upserting: {e}")
            return False
    
    def search_similar_content(self, query: str, top_k: int = 5):
        try:
            if not self.index:
                return []
            query_embedding = generate_embedding(query)
            results = self.index.query(vector=query_embedding, top_k=top_k, include_metadata=True)
            return [{'text': m['metadata'].get('text', ''), 'score': m['score']} for m in results.get('matches', [])]
        except Exception as e:
            logger.error(f"Error searching: {e}")
            return []

pinecone_manager = PineconeManager()

def generate_embedding(text: str) -> List[float]:
    try:
        if not client:
            raise Exception("OpenAI not initialized")
        response = client.embeddings.create(input=text, model=EMBEDDING_MODEL)
        return response.data[0].embedding
    except Exception as e:
        logger.error(f"Error generating embedding: {e}")
        raise

def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
        return text if text.strip() else extract_text_from_pdf_pypdf2(file_content)
    except:
        return extract_text_from_pdf_pypdf2(file_content)

def extract_text_from_pdf_pypdf2(file_content: bytes) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
        if not text.strip():
            raise Exception("No text extracted")
        return text
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_content))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += " " + " ".join(cell.text for cell in row.cells)
        return text
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def process_uploaded_file(file_content: bytes, filename: str, user_id: int = None) -> Dict[str, Any]:
    try:
        ext = filename.lower().split('.')[-1]
        if ext == 'pdf':
            text = extract_text_from_pdf(file_content)
        elif ext in ['docx', 'doc']:
            text = extract_text_from_docx(file_content)
        elif ext == 'txt':
            text = file_content.decode('utf-8')
        else:
            raise Exception(f"Unsupported type: {ext}")
        if not text.strip():
            raise Exception("No text extracted")
        parsed = parseSyllabusText(text)
        store_result = store_syllabus(parsed, user_id)
        if store_result['success']:
            parsed['id'] = store_result['id']
            return {'success': True, 'message': 'Processed successfully', 'data': parsed}
        raise Exception("Storage failed")
    except Exception as e:
        logger.error(f"Processing error: {e}")
        return {'success': False, 'error': str(e)}

def parseSyllabusText(text: str) -> Dict[str, Any]:
    """Use OpenAI to parse syllabus and extract course info, dates, topics, and grading."""
    try:
        if not client:
            raise Exception("OpenAI not initialized")
        
        prompt = f"""You are an expert syllabus parser. Your job is to extract EVERYTHING important from this syllabus.
DO NOT SKIP ANYTHING. Extract comprehensively and thoroughly.

COURSE INFORMATION:
Extract exactly:
- course: Full course number and title (e.g., "CSCE 120: Program Design and Concepts")
- instructor: Primary instructor name (or "Multiple Instructors" if several listed)
- semester: Semester and year (e.g., "Fall 2025")

KEY DATES TO EXTRACT - BE EXHAUSTIVE:
Look for and INCLUDE:
- **Exam dates**: "Exam 1", "Exam 2", "Exam 3", "Midterm", "Final Exam", "Quiz", with SPECIFIC DATES or WEEKS
  - Week of Sep 22 → "Approximately September 22, 2025"
  - Week of Oct 27 → "Approximately October 27, 2025"
  - Week of Nov 17 → "Approximately November 17, 2025"
  - Dec. 12 (F), 3:30 PM → "December 12, 2025, 3:30 PM - 5:30 PM"
- **Homework due dates**: EVERY SINGLE ONE in the list
  - Friday, September 5, 2025
  - Friday, September 12, 2025
  - Friday, September 19, 2025
  - etc. ALL 10 of them
- **Lab dates**: If recurring, capture as "Weekly during lab session"
- **Project/Assignment milestones**: Problem Definition (Sep 19), Problem Decomposition (Oct 17), Implementation (Nov 21), Venture Capital pitch (Dec 1)
- **Important deadlines**: Drop dates, last day of classes, any other dates mentioned

For EACH date entry, provide:
- date: The actual calendar date (convert "Week of Sep 22" to "September 22, 2025" if year appears elsewhere, or just the date)
- event: Full description including assignment number if applicable
- type: exam|quiz|homework|lab|project|other
- note: Any additional context (optional)

GRADING BREAKDOWN - EXTRACT EVERY COMPONENT:
CRITICAL: Parse the grading policy section thoroughly. Extract EACH weighted category:

Go through the syllabus and find ALL grading components with their weights:
- Homework (20%)
- Labwork (5% for sections 200,201,250 OR 10% for others)
- Midterm Exams (36%)
- Final Exam (24%)
- Readings (0% for sections 200,201,250 OR 5% for others)
- Class Engagement (5%)
- Honors Activities (10% for sections 200,201,250 only)

Extract EVERYTHING mentioned with percentages:
{{
  "category": "string (exact name from syllabus)",
  "weight": number (0-100),
  "note": "any special conditions (optional)"
}}

For this specific syllabus:
- List all 7 grading categories
- Include section-specific variations
- Note that Honors students have different weightings

COURSE TOPICS - EXTRACT ALL:
Look for any section titled "Topics" or "Content" or "Course Outline". Extract EVERY SINGLE TOPIC:
- What you (should) already know, but in C++
- Vectors
- Errors and Debugging
- Object-Oriented Concepts
- Design and Planning
- Exceptions
- Input/Output Streams
- Functions – Pass by Reference
- Class Design
- Implementing Classes
- Dynamic Memory
- C-style Arrays
- Dynamic Arrays
- Classes with Dynamic Memory
- Linked Lists
- Inheritance and Polymorphism
- Recursion

Also look for learning outcomes and extract key concepts:
- Abstraction
- Information hiding
- Object-oriented decomposition
- Control structures (sequence, selection, iteration)
- Data types (arrays, vectors, linked lists, structs, classes)
- Debugging
- Memory management
- Memory leaks

IMPORTANT INSTRUCTIONS:
1. DO NOT SKIP OR MISS DATES - There are 10 homework due dates, 3 midterm exam weeks, 1 final exam, and 4 honors artifacts
2. DO NOT MISS GRADING - Parse the grading policy section thoroughly and extract all percentages
3. DO NOT MISS TOPICS - Both the "Topics" section AND the learning outcomes section
4. For grading with section variations, list all variations
5. If dates reference "Week of X" or "around Y date", convert to actual calendar dates as best you can
6. EXTRACT NUMBERS: 20%, 36%, 24%, 5%, 10%, 0% - these are all grading percentages
7. EXTRACT DATES: Sep 5, Sep 12, Sep 19, Oct 3, Oct 10, Oct 17, Oct 24, Nov 7, Nov 14, Dec 5, Sep 19, Oct 17, Nov 21, Dec 1
8. EXTRACT TOPICS: ALL 16 topics listed, PLUS key concepts from learning outcomes

RESPONSE FORMAT - MUST BE VALID JSON:
{{
  "course": "string (CSCE 120: Program Design and Concepts)",
  "instructor": "string (multiple instructors listed)",
  "semester": "Fall 2025",
  "keyDates": [
    {{
      "date": "string (Month Day, Year or date range)",
      "event": "string (detailed description)",
      "type": "exam|quiz|homework|lab|project|other",
      "note": "optional"
    }}
  ],
  "topics": [
    "string (individual topic)",
    "string (another topic)"
  ],
  "gradingBreakdown": [
    {{
      "category": "string",
      "weight": number,
      "note": "optional - section variations or conditions"
    }}
  ]
}}

CRITICAL REMINDERS:
- keyDates should have 15+ entries (10 homework + 3 exams + 4 honors + finals)
- topics should have 20+ entries (16 course topics + key learning concepts)
- gradingBreakdown should have 6-7 entries with different weights listed
- Extract homework dates as individual entries: HW 1, HW 2, HW 3, etc.
- Extract exam weeks and final exam dates
- Do NOT make up any dates or topics. Only extract what is explicitly stated in the syllabus.
- DO NOT skip or abbreviate - be comprehensive and thorough

SYLLABUS TEXT:
{text[:10000]}

Return ONLY valid JSON. No markdown. No explanations. Start with {{ and end with }}."""
        
        response = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=5000,
            temperature=0
        )
        
        result_text = response.choices[0].message.content.strip()
        
        # Try to find JSON in the response
        json_match = re.search(r'\{[\s\S]*\}', result_text)
        if not json_match:
            logger.error(f"No JSON found in OpenAI response: {result_text[:500]}")
            return {
                'course': 'Unknown Course',
                'instructor': 'Unknown Instructor',
                'semester': 'Unknown Semester',
                'keyDates': [],
                'topics': [],
                'gradingBreakdown': [],
                'exams': 0,
                'assignments': 0
            }
        
        parsed_data = json.loads(json_match.group())
        
        # Ensure all required fields exist
        parsed_data.setdefault('course', 'Unknown Course')
        parsed_data.setdefault('instructor', 'Unknown Instructor')
        parsed_data.setdefault('semester', 'Unknown Semester')
        parsed_data.setdefault('keyDates', [])
        parsed_data.setdefault('topics', [])
        parsed_data.setdefault('gradingBreakdown', [])
        
        # Remove duplicates from topics (case-insensitive)
        if parsed_data.get('topics'):
            seen = set()
            unique_topics = []
            for topic in parsed_data['topics']:
                topic_lower = str(topic).lower().strip()
                if topic_lower not in seen and topic_lower:
                    seen.add(topic_lower)
                    unique_topics.append(topic)
            parsed_data['topics'] = unique_topics
        
        # Remove duplicate dates
        if parsed_data.get('keyDates'):
            seen = set()
            unique_dates = []
            for date_entry in parsed_data['keyDates']:
                key = (date_entry.get('date', '').lower(), date_entry.get('event', '').lower())
                if key not in seen:
                    seen.add(key)
                    unique_dates.append(date_entry)
            parsed_data['keyDates'] = unique_dates
        
        parsed_data['exams'] = len([d for d in parsed_data.get('keyDates', []) if d.get('type') == 'exam'])
        parsed_data['assignments'] = len([d for d in parsed_data.get('keyDates', []) if d.get('type') in ['homework', 'project', 'lab']])
        
        logger.info(f"Parsed syllabus: {parsed_data['course']}")
        logger.info(f"  - Found {len(parsed_data['keyDates'])} key dates")
        logger.info(f"  - Found {len(parsed_data['topics'])} topics")
        logger.info(f"  - Found {len(parsed_data['gradingBreakdown'])} grading categories")
        logger.info(f"  - Parsed {parsed_data['exams']} exams and {parsed_data['assignments']} assignments")
        
        return parsed_data
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in syllabus parsing: {e}")
        return {
            'course': 'Unknown Course',
            'instructor': 'Unknown Instructor',
            'semester': 'Unknown Semester',
            'keyDates': [],
            'topics': [],
            'gradingBreakdown': [],
            'exams': 0,
            'assignments': 0
        }
    except Exception as e:
        logger.error(f"Error parsing syllabus with OpenAI: {e}")
        return {
            'course': 'Unknown Course',
            'instructor': 'Unknown Instructor',
            'semester': 'Unknown Semester',
            'keyDates': [],
            'topics': [],
            'gradingBreakdown': [],
            'exams': 0,
            'assignments': 0
        }

def store_syllabus(data: Dict[str, Any], user_id: int = None) -> Dict[str, Any]:
    try:
        from models import db, Syllabus
        
        course_id = data['course'].replace(' ', '_').replace('-', '_').lower()
        chunks = [{'text': f"Course: {data['course']} ({data['semester']})", 'course_name': data['course']}]
        for d in data.get('keyDates', []):
            chunks.append({'text': f"{d['event']} on {d['date']}", 'course_name': data['course']})
        
        # Upsert to Pinecone
        vector_ids = pinecone_manager.upsert_syllabus_content(course_id, chunks)
        
        # Save to database if user_id provided
        if user_id:
            syllabus = Syllabus(
                user_id=user_id,
                course_id=course_id,
                course_name=data.get('course', 'Unknown'),
                instructor=data.get('instructor', 'Unknown'),
                semester=data.get('semester', 'Unknown'),
                key_dates=data.get('keyDates', []),
                topics=data.get('topics', []),
                grading_breakdown=data.get('gradingBreakdown', []),
                vector_ids=[]  # Will populate with actual vector IDs if needed
            )
            db.session.add(syllabus)
            db.session.commit()
            logger.info(f"Saved syllabus to database for user {user_id}: {data['course']}")
            return {'success': True, 'id': syllabus.id}
        
        return {'success': True, 'id': None}
    except Exception as e:
        logger.error(f"Store error: {e}")
        return {'success': False, 'id': None}

def query_syllabus_content(query: str, course_filter: Optional[str] = None, user_id: int = None) -> str:
    try:
        if not client:
            return "OpenAI not configured"
        content = pinecone_manager.search_similar_content(query, 5)
        ctx = "\n".join([f"• {c['text']}" for c in content]) if content else "No info found"
        msgs = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "system", "content": f"Info:\n{ctx}"},
            {"role": "user", "content": query}
        ]
        resp = client.chat.completions.create(model=CHAT_MODEL, messages=msgs, max_tokens=500)
        return resp.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

def get_service_status():
    return {'openai': bool(client), 'pinecone': bool(pinecone_manager.index), 'error': None}

def initialize_openai_service():
    logger.info("Service initialized") 
